"""Safe, best-effort document text extraction for student uploads."""

from __future__ import annotations

import asyncio
from dataclasses import dataclass
from io import BytesIO
from typing import Literal

from app.workspace.models import (
    EXTRACTED_TEXT_MAX_LENGTH,
    DocumentCreate,
    DocumentTextStatus,
    DocumentUpload,
    WorkspaceValidationError,
)

ExtractionKind = Literal["pdf", "docx", "text"]
UploadKind = ExtractionKind | Literal["image"]

#: Bound for pypdf/python-docx parsing, which runs off the event loop via
#: asyncio.to_thread. A crafted small PDF/DOCX can decompress to gigabytes or
#: pathologically stall the shared thread pool (decompression-bomb DoS); this
#: is the fallback used when a caller doesn't pass its own Settings-derived
#: value (see Settings.document_extraction_timeout_s).
DEFAULT_EXTRACTION_TIMEOUT_S = 8.0


@dataclass(frozen=True)
class _UploadFormat:
    mime: str
    kind: UploadKind
    allowed_mimes: frozenset[str]


_GENERIC_UPLOAD_MIME = "application/octet-stream"
_UPLOAD_FORMATS: dict[str, _UploadFormat] = {
    ".pdf": _UploadFormat("application/pdf", "pdf", frozenset({"application/pdf"})),
    ".docx": _UploadFormat(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
        frozenset({"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}),
    ),
    ".txt": _UploadFormat("text/plain", "text", frozenset({"text/plain"})),
    ".md": _UploadFormat(
        "text/markdown", "text", frozenset({"text/markdown", "text/x-markdown", "text/plain"})
    ),
    ".markdown": _UploadFormat(
        "text/markdown", "text", frozenset({"text/markdown", "text/x-markdown", "text/plain"})
    ),
    ".jpg": _UploadFormat("image/jpeg", "image", frozenset({"image/jpeg"})),
    ".jpeg": _UploadFormat("image/jpeg", "image", frozenset({"image/jpeg"})),
    ".png": _UploadFormat("image/png", "image", frozenset({"image/png"})),
    ".gif": _UploadFormat("image/gif", "image", frozenset({"image/gif"})),
    ".webp": _UploadFormat("image/webp", "image", frozenset({"image/webp"})),
    ".heic": _UploadFormat("image/heic", "image", frozenset({"image/heic"})),
    ".heif": _UploadFormat("image/heif", "image", frozenset({"image/heif"})),
    ".tif": _UploadFormat("image/tiff", "image", frozenset({"image/tiff"})),
    ".tiff": _UploadFormat("image/tiff", "image", frozenset({"image/tiff"})),
    ".bmp": _UploadFormat("image/bmp", "image", frozenset({"image/bmp"})),
    ".avif": _UploadFormat("image/avif", "image", frozenset({"image/avif"})),
}
_SUPPORTED_UPLOAD_EXTENSIONS = ", ".join(sorted(_UPLOAD_FORMATS))


async def prepare_document_upload(
    data: DocumentUpload, *, extraction_timeout_s: float = DEFAULT_EXTRACTION_TIMEOUT_S
) -> DocumentCreate:
    """Validate, normalize, and extract a student upload without blocking the event loop.

    The original bytes are retained only after format validation. Extraction
    failures after validation produce a truthful ``failed`` text status, never
    usable text. Both parsing passes are time-bounded (``extraction_timeout_s``)
    so a decompression-bomb-style PDF/DOCX can't stall the shared thread pool.
    """
    upload_format = _validate_upload_metadata(data)
    if upload_format.kind == "image":
        return _document_create(data, upload_format.mime, "unsupported", None)

    try:
        await asyncio.wait_for(
            asyncio.to_thread(_validate_extractable_content, upload_format.kind, data.content),
            timeout=extraction_timeout_s,
        )
    except TimeoutError as exc:
        labels = {"pdf": "PDF", "docx": "DOCX", "text": "UTF-8 text"}
        raise WorkspaceValidationError(
            f"document content is not a valid {labels[upload_format.kind]} file"
        ) from exc

    try:
        extracted_text = await asyncio.wait_for(
            asyncio.to_thread(_extract_text, upload_format.kind, data.content),
            timeout=extraction_timeout_s,
        )
    except TimeoutError:
        # Extraction is deliberately degradable (see _extract_text) — a timeout
        # is just another unreadable-file outcome, never an unhandled crash.
        extracted_text = None

    if not _has_usable_text(extracted_text):
        return _document_create(data, upload_format.mime, "failed", None)
    assert extracted_text is not None  # narrowed by _has_usable_text above
    return _document_create(
        data, upload_format.mime, "extracted", extracted_text[:EXTRACTED_TEXT_MAX_LENGTH]
    )


def _validate_upload_metadata(data: DocumentUpload) -> _UploadFormat:
    if _unsafe_filename(data.filename):
        raise WorkspaceValidationError("document filename must be a single safe filename")

    suffix = _filename_suffix(data.filename)
    upload_format = _UPLOAD_FORMATS.get(suffix)
    if upload_format is None:
        raise WorkspaceValidationError(
            f"unsupported document type; upload one of: {_SUPPORTED_UPLOAD_EXTENSIONS}"
        )

    declared_mime = data.mime.split(";", 1)[0].strip().lower()
    if declared_mime not in {*upload_format.allowed_mimes, _GENERIC_UPLOAD_MIME}:
        raise WorkspaceValidationError("document content type does not match its filename")
    if upload_format.kind == "image" and not _matches_image_signature(
        upload_format.mime, data.content
    ):
        raise WorkspaceValidationError("document image content does not match its filename")
    return upload_format


def _validate_extractable_content(kind: ExtractionKind, content: bytes) -> None:
    """Reject spoofed binary uploads before their bytes can be persisted.

    Plain text has no reliable file signature, so its intentionally simple
    contract is UTF-8 (with an optional BOM). PDF and DOCX must additionally
    be parseable by the same readers used for extraction.
    """
    try:
        if kind == "pdf":
            _validate_pdf(content)
        elif kind == "docx":
            _validate_docx(content)
        else:
            content.decode("utf-8-sig")
    except Exception as exc:
        labels = {"pdf": "PDF", "docx": "DOCX", "text": "UTF-8 text"}
        raise WorkspaceValidationError(
            f"document content is not a valid {labels[kind]} file"
        ) from exc


def _validate_pdf(content: bytes) -> None:
    if not content.startswith(b"%PDF-"):
        raise ValueError("missing PDF header")
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content), strict=False)
    if not reader.is_encrypted:
        len(reader.pages)


def _validate_docx(content: bytes) -> None:
    if not content.startswith(b"PK\x03\x04"):
        raise ValueError("missing ZIP header")
    from docx import Document as DocxDocument

    DocxDocument(BytesIO(content))


def _unsafe_filename(filename: str) -> bool:
    return (
        filename in {".", ".."}
        or "/" in filename
        or "\\" in filename
        or any(character in filename for character in ("\x00", "\r", "\n"))
    )


def _filename_suffix(filename: str) -> str:
    return f".{filename.rsplit('.', 1)[1].lower()}" if "." in filename else ""


def _matches_image_signature(mime: str, content: bytes) -> bool:
    signatures = {
        "image/jpeg": content.startswith(b"\xff\xd8\xff"),
        "image/png": content.startswith(b"\x89PNG\r\n\x1a\n"),
        "image/gif": content.startswith((b"GIF87a", b"GIF89a")),
        "image/webp": len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP",
        "image/tiff": content.startswith((b"II*\x00", b"MM\x00*")),
        "image/bmp": content.startswith(b"BM"),
    }
    if mime in signatures:
        return signatures[mime]
    return _matches_iso_base_media_image(mime, content)


def _matches_iso_base_media_image(mime: str, content: bytes) -> bool:
    if len(content) < 12 or content[4:8] != b"ftyp":
        return False
    brands = {content[index : index + 4] for index in range(8, len(content) - 3, 4)}
    if mime == "image/avif":
        return bool(brands & {b"avif", b"avis"})
    if mime == "image/heic":
        return bool(brands & {b"heic", b"heix", b"hevc", b"hevx", b"heim", b"heis"})
    return bool(brands & {b"mif1", b"msf1", b"heif"})


def _extract_text(kind: ExtractionKind, content: bytes) -> str | None:
    try:
        if kind == "pdf":
            return _extract_pdf(content)
        if kind == "docx":
            return _extract_docx(content)
        return content.decode("utf-8-sig")
    except Exception:
        # Extraction is deliberately degradable: valid but encrypted, scanned,
        # or otherwise unreadable files remain stored without usable text.
        return None


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content), strict=False)
    if reader.is_encrypted and reader.decrypt("") == 0:
        raise ValueError("encrypted PDF")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(BytesIO(content))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(
        "\t".join(cell.text for cell in row.cells)
        for table in document.tables
        for row in table.rows
    )
    return "\n".join(blocks)


def _document_create(
    data: DocumentUpload,
    mime: str,
    text_status: DocumentTextStatus,
    extracted_text: str | None,
) -> DocumentCreate:
    return DocumentCreate(
        title=data.title,
        doc_type=data.doc_type,
        filename=data.filename,
        mime=mime,
        content=data.content,
        text_status=text_status,
        extracted_text=extracted_text,
    )


def _has_usable_text(text: str | None) -> bool:
    return bool(
        text
        and any(character.isprintable() and not character.isspace() for character in text)
    )
